from docx import Document
#pip install python-docx
from pyvi import ViTokenizer, ViPosTagger
#pip install pyvi
document = Document('522H0148.docx')

paragraph1 = document.paragraphs[0]
paragraph2 = document.paragraphs[1]
paragraph3 = document.paragraphs[2]

if paragraph1.runs[0].font.name == 'Times New Roman' and paragraph1.runs[0].font.size == 914400:
    print('Dòng chữ 1: TỔNG LIÊN ĐOÀN LAO ĐỘNG VIỆT NAM - Font: Normal, Size: 14')
else:
    print('Dòng chữ 1 không đúng yêu cầu.')

if paragraph2.runs[0].font.bold and paragraph2.runs[0].font.size == 914400:
    print('Dòng chữ 2: TRƯỜNG ĐẠI HỌC TÔN ĐỨC THẮNG - Font: Bold, Size: 14')
else:
    print('Dòng chữ 2 không đúng yêu cầu.')

if paragraph3.runs[0].font.bold and paragraph3.runs[0].font.size == 914400:
    print('Dòng chữ 3: KHOA CÔNG NGHỆ THÔNG TIN - Font: Bold, Size: 14')
else:
    print('Dòng chữ 3 không đúng yêu cầu.')

# Kiểm tra lỗi chính tả
for paragraph in document.paragraphs:
    text = paragraph.text
    tokens = ViTokenizer.tokenize(text)
    pos_tags = ViPosTagger.postagging(tokens)
    for i in range(len(pos_tags[1])):
        if pos_tags[1][i] == 'Np' or pos_tags[1][i] == 'Nc':
            word = pos_tags[0][i]
            if not ViPosTagger.is_in_vn_dict(word):
                print(f'Lỗi chính tả tại dòng "{text}"')